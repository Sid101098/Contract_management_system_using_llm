import os
import PyPDF2
import docx
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.embeddings import OpenAIEmbeddings
from langchain.vectorstores import Chroma
from langchain.schema import Document
import logging
from datetime import datetime
import re

logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')

class DocumentProcessor:
    def __init__(self, persist_directory="./chroma_db"):
        self.embeddings = OpenAIEmbeddings()
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=1000,
            chunk_overlap=200,
            length_function=len
        )
        self.persist_directory = persist_directory
        self.vectorstore = None
        
    def load_documents(self, directory_path):
        """Load and process documents from directory"""
        documents = []
        
        for filename in os.listdir(directory_path):
            if filename.endswith('.pdf'):
                docs = self._process_pdf(os.path.join(directory_path, filename))
                documents.extend(docs)
            elif filename.endswith('.docx'):
                docs = self._process_docx(os.path.join(directory_path, filename))
                documents.extend(docs)
            elif filename.endswith('.txt'):
                docs = self._process_txt(os.path.join(directory_path, filename))
                documents.extend(docs)
        
        return documents
    
    def _process_pdf(self, file_path):
        """Extract text from PDF files"""
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                text = ""
                for page_num, page in enumerate(pdf_reader.pages):
                    text += f"\n--- Page {page_num + 1} ---\n"
                    text += page.extract_text()
                
                return [Document(
                    page_content=text,
                    metadata={
                        "source": os.path.basename(file_path),
                        "file_type": "pdf",
                        "processed_date": datetime.now().isoformat()
                    }
                )]
        except Exception as e:
            logging.error(f"Error processing PDF {file_path}: {e}")
            return []
    
    def _process_docx(self, file_path):
        """Extract text from DOCX files"""
        try:
            doc = docx.Document(file_path)
            text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
            
            return [Document(
                page_content=text,
                metadata={
                    "source": os.path.basename(file_path),
                    "file_type": "docx",
                    "processed_date": datetime.now().isoformat()
                }
            )]
        except Exception as e:
            logging.error(f"Error processing DOCX {file_path}: {e}")
            return []
    
    def _process_txt(self, file_path):
        """Extract text from TXT files"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                text = file.read()
            
            return [Document(
                page_content=text,
                metadata={
                    "source": os.path.basename(file_path),
                    "file_type": "txt",
                    "processed_date": datetime.now().isoformat()
                }
            )]
        except Exception as e:
            logging.error(f"Error processing TXT {file_path}: {e}")
            return []
    
    def chunk_documents(self, documents):
        """Split documents into chunks"""
        return self.text_splitter.split_documents(documents)
    
    def create_vectorstore(self, documents):
        """Create and persist vector store"""
        try:
            self.vectorstore = Chroma.from_documents(
                documents=documents,
                embedding=self.embeddings,
                persist_directory=self.persist_directory
            )
            self.vectorstore.persist()
            logging.info("Vector store created and persisted successfully")
            return self.vectorstore
        except Exception as e:
            logging.error(f"Error creating vector store: {e}")
            raise
    
    def load_vectorstore(self):
        """Load existing vector store"""
        try:
            self.vectorstore = Chroma(
                persist_directory=self.persist_directory,
                embedding_function=self.embeddings
            )
            return self.vectorstore
        except Exception as e:
            logging.error(f"Error loading vector store: {e}")
            return None
